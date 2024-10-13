from script.utils import chunks
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
from docx.shared import Pt
import pandas as pd
import numpy as np
import os
import shutil
from tqdm import tqdm
import glob
from zipfile import ZipFile 
import concurrent
pd.options.mode.chained_assignment = None 
import json

class MailMerge():
    def __init__(self, dataset ,label, file_per_zip = 100, parallel = False):
        self.dataset = dataset
        self.label = label
        self.file_per_zip = file_per_zip
        self.multicore = parallel

    def safe_cast(val, to_type, default=None):
        try:
            return to_type(val)
        except (ValueError, TypeError):
            return default
            
    def row_to_pdf(self, data):
        for i in data.keys():
            if str(data.get(i)) == 'nan':
                data[i] = ''
            else:
                data[i] = str(data[i])

        
        f = open('template.docx', 'rb')
        doc = Document(f)

        rep_zero = ['card_no', 'consult_id', 'claim_id','claim_id_rx','consult_fee', 'order_id' ,
                   'deliv_coverage_by_insurance', 'total_consult_+_rx',
                    'total', 'rx_excess', 'excess_consult', 'excess', 'amag_discount']

        for i in rep_zero:
            data[i] = data[i].replace('.0', '')
        

        data['datetime'] = pd.to_datetime(data['date']).strftime('%d/%m/%Y') 
        
        ## General Information
        doc.tables[0].cell(0, 1).text = data['doctor_name']
        doc.tables[0].cell(1, 1).text = data['doctor_department']
        doc.tables[0].cell(2, 1).text = str(data['doctor_sip'])
        doc.tables[0].cell(3, 1).text = str(data['doctor_str'])
        doc.tables[0].cell(0, 3).text = data['name']
        doc.tables[0].cell(1, 3).text = data['gender']
        doc.tables[0].cell(2, 3).text = str(data['dob'])
        doc.tables[0].cell(3, 3).text = data['card_no']
        doc.tables[0].cell(4, 3).text = data['payor']
        doc.tables[0].cell(5, 3).text = data['corporate']

        for i in range(6):
            if i < 4:
                doc.tables[0].cell(i, 1).paragraphs[0].runs[0].font.size = Pt(8)
            doc.tables[0].cell(i, 3).paragraphs[0].runs[0].font.size = Pt(8)
        
        ## complaint & diagnosis
        doc.paragraphs[2].text = data['chief_complaints']
        doc.paragraphs[2].runs[0].font.size= Pt(8)
        doc.paragraphs[4].text = str(data['diagnosis']) if str(data['diagnosis']) != '' else ''
        doc.paragraphs[4].runs[0].font.size= Pt(8)
        
        doc.paragraphs[6].text = data['icdx']
        doc.paragraphs[6].runs[0].font.size= Pt(8)

        doc.paragraphs[8].text = str(data['suggestion']) if str(data['suggestion']) != '' else '' 
        doc.paragraphs[8].runs[0].font.size= Pt(8)
        
        ## Consultation Detail
        doc.tables[1].cell(1, 1).text = str(data['consult_id'])
        doc.tables[1].cell(1, 2).text = str(data['claim_id'])
        doc.tables[1].cell(2, 2).text = str(data['claim_id_rx'])
        doc.tables[1].cell(1, 3).text = data['datetime']
        doc.tables[1].cell(1, 4).text = 'IDR '+ f'{float(data["consult_fee"]):,}'.replace('.0', '')
        doc.tables[1].cell(1, 7).text = 'IDR '+ f'{float(data["consult_fee"]):,}'.replace('.0', '')
        doc.tables[1].cell(1, 4).paragraphs[0].alignment = 2
        doc.tables[1].cell(1, 7).paragraphs[0].alignment = 2
        doc.tables[1].cell(2, 1).text = str(data['order_id'])
        doc.tables[1].cell(2, 3).text = '' if str(data['order_created_date']).lower() in ('nat', 'nan', '') else pd.to_datetime(data['order_created_date'], errors='coerce').strftime('%d/%m/%Y') + ' ' + str(data['order_created_time'])
        
        
        doc.tables[1].cell(1, 1).paragraphs[0].runs[0].font.size= Pt(8)
        doc.tables[1].cell(1, 2).paragraphs[0].runs[0].font.size= Pt(8)
        doc.tables[1].cell(2, 2).paragraphs[0].runs[0].font.size= Pt(8)
        doc.tables[1].cell(1, 3).paragraphs[0].runs[0].font.size= Pt(8)
        doc.tables[1].cell(1, 4).paragraphs[0].runs[0].font.size= Pt(8)
        doc.tables[1].cell(1, 7).paragraphs[0].runs[0].font.size= Pt(8)
        doc.tables[1].cell(2, 1).paragraphs[0].runs[0].font.size= Pt(8)
        doc.tables[1].cell(2, 3).paragraphs[0].runs[0].font.size= Pt(8)


        for i in range(1,13):
            j = i + 2 #f'{value:,}' 
            if str(data[f'obat_{i}']) != '':
                doc.tables[1].cell(j, 0).text = data[f'obat_{i}']
                # doc.tables[1].cell(j, 0).paragraphs[0].alignment = 2
                harga_i = float(data[f'harga_{i}'])
                total_i = float(data[f'total_{i}'])
                doc.tables[1].cell(j, 4).text = 'IDR '+ f'{harga_i:,}'.replace('.0', '')
                doc.tables[1].cell(j, 4).paragraphs[0].alignment = 2
                doc.tables[1].cell(j, 5).text = str(data[f'jumlah_{i}']).replace('.0', '')
                doc.tables[1].cell(j, 6).text = data[f'unit_obat_{i}']
                doc.tables[1].cell(j, 7).text = 'IDR '+ f'{total_i:,}'.replace('.0', '')
                doc.tables[1].cell(j, 7).paragraphs[0].alignment = 2
        
                doc.tables[1].cell(j, 6).paragraphs[0].runs[0].italic = True
                
                doc.tables[1].cell(j, 0).paragraphs[0].runs[0].font.size= Pt(8)
                doc.tables[1].cell(j, 4).paragraphs[0].runs[0].font.size= Pt(8)
                doc.tables[1].cell(j, 5).paragraphs[0].runs[0].font.size= Pt(8)
                doc.tables[1].cell(j, 6).paragraphs[0].runs[0].font.size= Pt(8)
                doc.tables[1].cell(j, 7).paragraphs[0].runs[0].font.size= Pt(8)


        deliv =  '0' if data['deliv_coverage_by_insurance'] =='' else f'{float(data["deliv_coverage_by_insurance"]):,}'.replace('.0', '')
        total_all = f'{float(data["total"]):,}' #str(data['total'])
        total_ins = f'{float(data["total_consult_+_rx"]):,}' #str(data['total'])
        oop1 = float(data['total']) - float(data['total_consult_+_rx'])
        oop =f'{oop1:,}' #str(data['total'])

        doc.tables[1].cell(16, 4).text = 'IDR ' + deliv.replace('.0','')
        doc.tables[1].cell(16, 7).text = 'IDR ' + deliv.replace('.0','')
        doc.tables[1].cell(18, 7).text = 'IDR '+ total_all.replace('.0','')
        doc.tables[1].cell(19, 7).text = 'IDR '+ total_ins.replace('.0','')
        doc.tables[1].cell(20, 7).text = 'IDR '+ oop.replace('.0','')
        doc.tables[1].cell(16, 4).paragraphs[0].runs[0].font.size= Pt(8)
        doc.tables[1].cell(16, 7).paragraphs[0].runs[0].font.size= Pt(8)
        doc.tables[1].cell(18, 7).paragraphs[0].runs[0].font.size= Pt(8)
        doc.tables[1].cell(19, 7).paragraphs[0].runs[0].font.size= Pt(8)
        doc.tables[1].cell(20, 7).paragraphs[0].runs[0].font.size= Pt(8)
        doc.tables[1].cell(16, 4).paragraphs[0].alignment = 2
        doc.tables[1].cell(16, 7).paragraphs[0].alignment = 2
        doc.tables[1].cell(18, 7).paragraphs[0].alignment = 2
        doc.tables[1].cell(19, 7).paragraphs[0].alignment = 2
        doc.tables[1].cell(20, 7).paragraphs[0].alignment = 2
        doc.tables[1].cell(18, 7).paragraphs[0].runs[0].bold = True
        doc.tables[1].cell(19, 7).paragraphs[0].runs[0].bold = True
        doc.tables[1].cell(20, 7).paragraphs[0].runs[0].bold = True
        name = data['name'].replace('/', ' ').title()
        consult_on = pd.to_datetime(data['date'], errors = 'coerce').strftime('%Y%m%d') #+'_'+str(data['time']).replace(':', '')
        consult_id = str(data['consult_id']).replace('.0', '')
        agg = data['aggregator_name']
        payor = data['payor']

        if not os.path.exists(f'.tmp/{self.label}/{agg}/{payor}'):
            os.makedirs(f'.tmp/{self.label}/{agg}/{payor}')
        
        tmp_name = f'.tmp/{self.label}/{agg}/{payor}/{consult_on}_{name}_Consultation_Receipt_{consult_id}.docx'
        doc.save(tmp_name)
        f.close()
        
        subprocess.run(['libreoffice', '--convert-to', 'pdf' ,
                        tmp_name, '--outdir', f'.tmp/{self.label}/{agg}/{payor}']
                       ,stdout=subprocess.DEVNULL,
                        stderr=subprocess.DEVNULL
                    )
        
        os.remove(tmp_name)


        data['claim_id_all'] = data['claim_id'] +' & '+ data['claim_id_rx']
        #data['date'] = pd.to_datetime(data['date']).strftime('%Y-%m-%d') 
        sel_dt = ['date','time', 'name', 'card_no', 'gdt_id', 'consult_id','order_id', 'order_id_2'
        ,'icdx', 'gpsp', 'claim_id_all', 'consult_fee', 'prescription_fee', 'total_consult_+_rx','excess',
        'total', 'payor', 'corporate','doctor_name','doctor_department','diagnosis','aggregator_name','excess',
                  'amag_discount', 'deliv_coverage_by_insurance', 'rx_excess', 'excess_consult', 'policy_id', 'member_id']

        
        recap = {x: data[x] for x in sel_dt}

        return recap

    def create_recap(self, recap, df_success_file):

        df_success_file_mrg =recap.merge(df_success_file, how = 'inner', on = 'consult_id')
        
        to_int = ['consult_fee', 'deliv_coverage_by_insurance', 'total_consult_+_rx',
                    'total', 'rx_excess', 'excess_consult', 'excess', 'amag_discount', 'prescription_fee']
        for i in to_int:
            df_success_file_mrg[i] = pd.to_numeric(df_success_file_mrg[i]).astype('Int64')
        
        map_recap = {
            'date': 'Tanggal', 'name': 'Nama Peserta', 'card_no': 'Nomor Kartu', 
            'gdt_id' : 'User ID', 'consult_id':'Consult ID', 'icdx': 'ICD', 'gpsp':'GP/SP',
            'claim_id_all': 'Klaim ID', 'consult_fee': 'Biaya Konsul', 'prescription_fee': 'Biaya Obat',
            'excess': 'Excess' ,'total':'Total Raw', 'total_consult_+_rx':'Total Net' , 
            'amag_discount':'Amag Discount','payor':'Payor', 'corporate': 'Nama Perusahaan',
            'order_id': 'Order ID 1','order_id_2': 'Order ID 2','doctor_name':'Nama Dokter',
            'doctor_department': 'Departemen Dokter', 'time':'Time','diagnosis':'Diagnosa', 
            'aggregator_name':'Aggregator', 'policy_id':'Polis ID', 'member_id':'Member ID', 
            'rx_excess':'Rx Excess','excess_consult': 'Excess Consult',
            'zip_file': 'List Invoice', 'pdf_file': 'PDF file',
            'deliv_coverage_by_insurance': 'Biaya Pengantaran', 
        }

        df_success_file_mrg['cat_recap'] = np.where((df_success_file_mrg['aggregator_name'] == 'Admedika') &
                                                (df_success_file_mrg['payor'] == 'GI'), 'admedika_gi',
                                        np.where((df_success_file_mrg['aggregator_name'] == 'Admedika') &
                                                (df_success_file_mrg['payor'] == 'AMAG'), 'admedika_amag',
                                        np.where(df_success_file_mrg['aggregator_name'] == 'Meditap', 'meditap',
                                                 'general' 
        )))
        
        df_success_file_mrg = df_success_file_mrg.rename(columns = map_recap)
        #print(df_success_file.head())
        df_success_file_mrg = df_success_file_mrg.sort_values(['Tanggal', 'Nama Peserta'])


        col_maps = {
            'general' : ['Tanggal', 'Nama Peserta', 'Nomor Kartu', 'User ID', 'Consult ID', 'ICD', 'GP/SP', 
                            'Klaim ID', 'Biaya Konsul', 'Biaya Obat', 'Total', 'List Invoice', 'Payor', 
                         'Nama Perusahaan', 'Order ID 1', 'Order ID 2', 'Nama Dokter', 'Departemen Dokter', 
                         'Time', 'Diagnosa', 'Aggregator', 'Polis ID', 'Member ID', 'Excess Consult', 'Rx Excess'],
            'admedika_gi' : ['Tanggal', 'Nama Peserta', 'Nomor Kartu', 'User ID', 'Consult ID', 'ICD', 'GP/SP', 
                            'Klaim ID', 'Biaya Konsul', 'Biaya Obat', 'Total', 
                               'Biaya Pengantaran', 'List Invoice', 'Payor', 
                         'Nama Perusahaan', 'Order ID 1', 'Order ID 2', 'Nama Dokter', 'Departemen Dokter', 
                         'Time', 'Diagnosa', 'Aggregator', 'Polis ID', 'Member ID', 'Excess Consult', 
                         'Rx Excess'],            
            'admedika_amag' : ['Tanggal', 'Nama Peserta', 'Nomor Kartu', 'User ID', 'Consult ID', 'ICD', 'GP/SP', 
                            'Klaim ID', 'Biaya Konsul', 'Biaya Obat', 'Total','Amag Discount', 'List Invoice', 'Payor', 
                         'Nama Perusahaan', 'Order ID 1', 'Order ID 2', 'Nama Dokter', 'Departemen Dokter', 
                         'Time', 'Diagnosa', 'Aggregator', 'Polis ID', 'Member ID', 'Excess Consult', 
                         'Rx Excess'], 
            'meditap': ['Tanggal', 'Nama Peserta', 'Nomor Kartu', 'User ID', 
                        'Consult ID', 'ICD', 'GP/SP', 'Klaim ID', 'Biaya Konsul', 
                        'Biaya Obat', 'Total', 'Excess' ,  'List Invoice', 'Payor', 
                         'Nama Perusahaan', 'Order ID 1', 'Order ID 2', 'Nama Dokter', 'Departemen Dokter', 
                         'Time', 'Diagnosa', 'Aggregator', 'Polis ID', 'Member ID', 'Excess Consult', 'Rx Excess']
            
            }
        
        for i in ['general', 'admedika_gi', 'admedika_amag', 'meditap']:
            df_success_agg = df_success_file_mrg[df_success_file_mrg['cat_recap'] == i]
            
            if i != 'meditap':
                df_success_agg['Biaya Konsul'] = (pd.to_numeric(df_success_agg['Biaya Konsul'], 
                                                    errors = 'coerce').fillna(0) -
                                                   pd.to_numeric(df_success_agg['Excess Consult'],
                                                                 errors = 'coerce').fillna(0)).astype('Int64')
                
                df_success_agg['Biaya Obat'] = (pd.to_numeric(df_success_agg['Biaya Obat'], 
                                                               errors = 'coerce').fillna(0) - 
                                                 pd.to_numeric(df_success_agg['Rx Excess'], 
                                                               errors = 'coerce').fillna(0)).astype('Int64')
                                                 
                df_success_agg['Total'] = df_success_agg['Total Net'].astype('Int64')
            else:
                df_success_agg['Total'] = df_success_agg['Total Net'].astype('Int64')

            df_success_agg['Tanggal'] = pd.to_datetime(df_success_agg['Tanggal']).dt.date
            df_success_agg[col_maps.get(i)].to_excel(f'output/{self.label}/recap_{i}.xlsx', index = False)
            
        
    def chunk_and_zip(self):   
        if os.path.exists(f'output/{self.label}.zip'):
            os.remove(f'output/{self.label}.zip')
            
        fls = glob.glob(f'.tmp/{self.label}/*/*/*')
        fls = sorted(fls)
        
        unique_agg = list(dict.fromkeys([i.split('/')[2] for i in fls]))
        unique_pay = list(dict.fromkeys([i.split('/')[3] for i in fls]))
        
        success_file = []
        
        for i in unique_agg:
            for j in unique_pay:
                ij = [x for x in fls if x.split('/')[2] == i and x.split('/')[3] == j]
                k = 1
                for files in chunks(ij, self.file_per_zip):
                    files = sorted(files)
                    with ZipFile(f'output/{self.label}/{i}_{j}_{k}.zip','w') as zip: 
                        for file_each in files:
                            zip.write(file_each, file_each.split('/')[-1]) 
                            success_file.append([os.path.basename(file_each),f'{i}_{j}_{k}.zip'])
                    k = k + 1

        
        success_file_consult = [x[0].split('_')[-1].replace('.pdf', '') for x in success_file]
        
        self.errors.to_excel(f'output/{self.label}/errors.xlsx', index = False)
        
        df_success_file = pd.DataFrame({'pdf_file': [x[0] for x in success_file],
                                        'zip_file': [x[1] for x in success_file],
                                       'consult_id': success_file_consult})

        self.create_recap(self.success, df_success_file)
        #self.dataset[['date', 'name', 'card_no', 'gdt_id', 
        #            'consult_id','icdx','gpsp','claim_id', 'consult_price','prescription_fee', 'total']]

        fls2 = glob.glob(f'output/{self.label}/*')
        
        with ZipFile(f'output/{self.label}.zip','w') as zip:
            for file_each in fls2:
                zip.write(file_each, file_each.split('/')[-1]) 
        shutil.rmtree(f'output/{self.label}')

    def run(self):
        if os.path.exists(f'.tmp/{self.label}'):
            shutil.rmtree(f'.tmp/{self.label}')

        success = []
        errors = []
        for index, row in tqdm(self.dataset.iterrows()):
            rcp = self.row_to_pdf(row)
            try:
                #rcp = self.row_to_pdf(row)
                success.append(rcp)
            except Exception as e:
                errors.append(row)
                print(e)

        self.errors = pd.DataFrame(errors)
        self.success = pd.DataFrame(success)
        self.chunk_and_zip()


def run_mail_merge(file, label, sheet_name = 'data', file_per_zip = 100):
    df = pd.read_excel(file)
    if not os.path.exists(f'output/{label}'):
        os.makedirs(f'output/{label}')
    else:
        #shutil.rmtree(f'output/{label}')
        raise Exception("Error: Anda sedang menjalankan mail merge dengan label yang sama, silahkan ganti salah satu label")
    MailMerge(df, label, file_per_zip = file_per_zip).run()
    #shutil.rmtree(f'output/{label}')
    
